"""Export endpoint — supports Markdown and DOCX output."""
from __future__ import annotations

import io
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.v1.auth import get_current_user
from app.core.database import get_db
from app.models.conversation import Message
from app.models.project import Project
from app.models.user import User

router = APIRouter(prefix="/export", tags=["export"])


class ExportRequest(BaseModel):
    project_id: str
    format: str = "markdown"  # "markdown" | "docx"
    conversation_id: str | None = None


@router.post("")
async def export_project(
    body: ExportRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
) -> Response:
    # Verify ownership
    proj_result = await db.execute(
        select(Project).where(
            Project.id == body.project_id,
            Project.owner_id == current_user.id,
        )
    )
    project = proj_result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    # Gather messages
    query = select(Message).join(Message.conversation)
    from app.models.conversation import Conversation  # noqa: PLC0415
    query = query.where(Conversation.project_id == body.project_id)
    if body.conversation_id:
        query = query.where(Message.conversation_id == body.conversation_id)
    query = query.order_by(Message.created_at.asc())

    msg_result = await db.execute(query)
    messages = msg_result.scalars().all()

    if body.format == "docx":
        return _export_docx(project, messages)
    return _export_markdown(project, messages)


def _export_markdown(project: Project, messages: list[Message]) -> Response:
    lines = [
        f"# {project.title}",
        f"",
        f"**Project:** {project.document_type.replace('_', ' ').title()}  ",
        f"**Citation Style:** {project.citation_style.upper()}  ",
        f"**Language:** {project.language.upper()}  ",
        f"**Exported:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}",
        "",
        "---",
        "",
    ]
    for msg in messages:
        if msg.role == "user":
            lines.append(f"**You:** {msg.content}")
        elif msg.role == "assistant":
            lines.append(f"\n**Skripsiku ({msg.mode_used or 'AI'}):**\n")
            lines.append(msg.content)
        lines.append("")

    content = "\n".join(lines)
    filename = f"skripsiku-{project.title[:30].replace(' ', '-')}.md"
    return Response(
        content=content.encode("utf-8"),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_docx(project: Project, messages: list[Message]) -> StreamingResponse:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise HTTPException(status_code=501, detail="DOCX export requires python-docx")

    doc = Document()
    doc.core_properties.title = project.title
    doc.core_properties.author = "Skripsiku"

    title = doc.add_heading(project.title, level=0)
    title.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    doc.add_paragraph(
        f"Document Type: {project.document_type.replace('_', ' ').title()} | "
        f"Citation: {project.citation_style.upper()} | "
        f"Exported: {datetime.utcnow().strftime('%Y-%m-%d')}"
    )
    doc.add_paragraph("")

    for msg in messages:
        if msg.role == "user":
            p = doc.add_paragraph()
            run = p.add_run("You: ")
            run.bold = True
            p.add_run(msg.content)
        elif msg.role == "assistant":
            p = doc.add_paragraph()
            run = p.add_run(f"Skripsiku ({msg.mode_used or 'AI'}): ")
            run.bold = True
            run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
            doc.add_paragraph(msg.content)
        doc.add_paragraph("")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"skripsiku-{project.title[:30].replace(' ', '-')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
